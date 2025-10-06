import selenium
from selenium import webdriver 
import openpyxl
import time
import random
import os
import docx

dl = os.listdir('/home/diablo/Desktop/scrpy/phd')
dr = webdriver.Chrome(r'/home/diablo/chromedriver')
for dx in dl:
	wb = openpyxl.load_workbook('/home/diablo/Desktop/scrpy/phd/'+dx)
	sh = wb.active
	c = sh.max_row
	i = 1
	
	for x in range(1, c+1):
		url = sh.cell(row=x,column=2).value
		try:
			dr.get(url)
		except:
			continue
		wxx = docx.Document()
		wxy = docx.Document()
		zx = dr.find_element_by_class_name('single_course_section')
		zc = zx.text
		tx = zx.find_elements_by_tag_name('table')
		for ty in tx:
			try:
				zc = zc.replace(ty.text, '\n')
			except:
				pass
			wxy.add_paragraph(ty.get_attribute('innerHTML'))
			wxy.save('/home/diablo/Desktop/scrpy/phd_/'+dx+'--'+sh.cell(row=x,column=1).value+'2.docx')
		
		try:
			ex1 = zx.find_element_by_class_name('bodyslot')
			zc = zc.replace(ex1.text, '\n')
		except:
			pass
		try:
			ex2 = zx.find_element_by_class_name('college-course-name')
			zc = zc.replace(ex2.text, '\n')
		except:
			pass
		try:
			ex3 = zx.find_element_by_class_name('more_links')
			zc = zc.replace(ex3.text, '\n')
		except:
			pass
		try:
			ex4 = zx.find_element_by_class_name('author-section')
			zc = zc.replace(ex4.text, '\n')
		except:
			pass
		try:
			ex5 = zx.find_element_by_class_name('college-news-link')
			zc = zc.replace(ex5.text, '\n')
		except:
			pass
		wxx.add_paragraph(zc)
		wxx.save('/home/diablo/Desktop/scrpy/phd_/'+dx+'--'+sh.cell(row=x,column=1).value+'1.docx')
		i += 1

		time.sleep(random.randint(10,20))
dr.quit()